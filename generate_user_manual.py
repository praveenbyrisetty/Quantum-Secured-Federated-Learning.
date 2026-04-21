"""
Generate FLQC User Manual - matching friend's format exactly.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def add_bottom_border(paragraph, color="8B0000", size="12"):
    """Add a colored bottom border line to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_top_border(paragraph, color="8B0000", size="12"):
    """Add a colored top border line to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def create():
    doc = Document()

    # ===== PAGE SETUP (A4) =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'Times New Roman'
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # ================================================================
    # HEADER: "Project Title | 2026" with dark red bottom border
    # ================================================================
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(2)

    # Project title
    run_title = header_para.add_run(
        "Quantum-Secured Federated Learning for Skin Lesion Classification"
    )
    run_title.font.size = Pt(16)
    run_title.font.name = 'Times New Roman'

    # Separator " | "
    run_sep = header_para.add_run("  ")
    run_sep.font.size = Pt(16)

    # Year
    run_year = header_para.add_run("2026")
    run_year.font.size = Pt(16)
    run_year.font.name = 'Times New Roman'
    run_year.font.color.rgb = RGBColor(100, 130, 160)

    # Dark red bottom border line
    add_bottom_border(header_para, color="8B0000", size="18")

    doc.add_paragraph()  # small gap

    # ================================================================
    # "User Manual" heading
    # ================================================================
    um = doc.add_paragraph()
    um.paragraph_format.space_after = Pt(8)
    run = um.add_run("User Manual")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    # ================================================================
    # STEPS
    # ================================================================

    # --- Step 1 ---
    s1 = doc.add_paragraph()
    s1.paragraph_format.space_before = Pt(6)
    s1.paragraph_format.space_after = Pt(4)
    r = s1.add_run("Step 1: Launch the Application")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_1 = [
        "Open a terminal or command prompt on your system",
        "Navigate to the FLQC project directory",
        'Activate the virtual environment using "venv\\Scripts\\activate" (Windows)',
        'Run the command: streamlit run server.py',
        "The application opens in your browser at http://localhost:8501",
    ]
    for b in bullets_1:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 2 ---
    s2 = doc.add_paragraph()
    s2.paragraph_format.space_before = Pt(6)
    s2.paragraph_format.space_after = Pt(4)
    r = s2.add_run("Step 2: Configure Training Settings")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_2 = [
        "Locate the sidebar on the left side of the interface",
        'Select an aggregation strategy from the dropdown (default: "Krum + Trimmed Mean")',
        "Set the number of training rounds using the slider (1 to 10)",
        "Verify the device status shown at the bottom (GPU or CPU)",
    ]
    for b in bullets_2:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 3 ---
    s3 = doc.add_paragraph()
    s3.paragraph_format.space_before = Pt(6)
    s3.paragraph_format.space_after = Pt(4)
    r = s3.add_run("Step 3: Start Federated Training")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_3 = [
        'Click the "START TRAINING" button in the sidebar',
        "The system initializes a global CNN model and creates 3 hospital clients",
        "Each hospital generates a quantum key using the E91 protocol (Qiskit simulation)",
        "The CHSH inequality test verifies entanglement (S > 2.0 required)",
        "Training begins with live progress displayed per hospital",
    ]
    for b in bullets_3:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 4 ---
    s4 = doc.add_paragraph()
    s4.paragraph_format.space_before = Pt(6)
    s4.paragraph_format.space_after = Pt(4)
    r = s4.add_run("Step 4: Monitor Security Reports")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_4 = [
        "After each round, a security report is displayed automatically",
        "The report shows anomaly detection results, CHSH verification values, and encryption status",
        "Clients with failed CHSH tests (S \u2264 2.0) are blocked from contributing",
        "Differential Privacy budget (\u03b5) tracking is shown per client and cumulatively",
    ]
    for b in bullets_4:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 5 ---
    s5 = doc.add_paragraph()
    s5.paragraph_format.space_before = Pt(6)
    s5.paragraph_format.space_after = Pt(4)
    r = s5.add_run("Step 5: View Training Results")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_5 = [
        "Once all rounds complete, training results are displayed with charts",
        "Loss and accuracy graphs show the model's learning progress over rounds",
        "A security summary confirms encryption, aggregation strategy, and DP guarantees",
        "The global model is automatically evaluated on the full HAM10000 test set",
    ]
    for b in bullets_5:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 6 ---
    s6 = doc.add_paragraph()
    s6.paragraph_format.space_before = Pt(6)
    s6.paragraph_format.space_after = Pt(4)
    r = s6.add_run("Step 6: Review Model Evaluation")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_6 = [
        "Check the test accuracy, test loss, and total test samples evaluated",
        "Expand the confusion matrix to see per-class predictions vs actual labels",
        "Review per-class accuracy bars for all 7 skin lesion types",
        "Compare performance across classes to identify strengths and weaknesses",
    ]
    for b in bullets_6:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 7 ---
    s7 = doc.add_paragraph()
    s7.paragraph_format.space_before = Pt(6)
    s7.paragraph_format.space_after = Pt(4)
    r = s7.add_run("Step 7: Test Skin Lesion Prediction")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_7 = [
        "Scroll down to the prediction section after training is complete",
        "Upload a dermoscopic skin lesion image (PNG, JPG, JPEG, BMP, or WebP)",
        "The system classifies the image into one of 7 HAM10000 classes",
        "View the predicted class, confidence score, and severity level",
        "Read the recommended medical precautions displayed for the predicted condition",
    ]
    for b in bullets_7:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # --- Step 8 ---
    s8 = doc.add_paragraph()
    s8.paragraph_format.space_before = Pt(6)
    s8.paragraph_format.space_after = Pt(4)
    r = s8.add_run("Step 8: Exit the System")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    bullets_8 = [
        "Once done, close the browser tab to exit the Streamlit interface",
        "Press Ctrl+C in the terminal to stop the Streamlit server",
        "Deactivate the virtual environment by typing \"deactivate\"",
    ]
    for b in bullets_8:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'

    # ================================================================
    # FOOTER: dark red top border + "Department of CSE, GMRIT    Page XX"
    # ================================================================
    # Add some spacing before footer
    doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.space_before = Pt(6)
    add_top_border(footer_para, color="8B0000", size="18")

    # Add tab stop for right-aligned page number
    tab_stops = footer_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)

    run_dept = footer_para.add_run("Department of CSE, GMRIT")
    run_dept.font.size = Pt(10)
    run_dept.font.name = 'Times New Roman'

    run_tab = footer_para.add_run("\t")

    run_page = footer_para.add_run("Page 62")
    run_page.font.size = Pt(10)
    run_page.font.name = 'Times New Roman'

    # ===== SAVE =====
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'FLQC_User_Manual_Final.docx')
    out = os.path.abspath(out)
    doc.save(out)
    print(f"[DONE] User Manual saved to: {out}")


if __name__ == "__main__":
    create()
